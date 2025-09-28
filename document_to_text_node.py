import logging
import base64
import os
import tempfile
from typing import Dict, Any, List, Optional
from datetime import datetime
import mimetypes
from pathlib import Path
from models.workflow import NodeResult

logger = logging.getLogger("workflow_api")

# Install these dependencies in requirements.txt:
# python-docx
# PyPDF2
# python-pptx
# textract

try:
    import PyPDF2
    from docx import Document
    import textract
    DEPENDENCIES_LOADED = True
except ImportError:
    logger.warning("Document conversion dependencies not installed. Some functionality may be limited.")
    DEPENDENCIES_LOADED = False

async def handle_document_to_text_node(
    node_id: str,
    node_data: Dict[str, Any],
    inputs: Dict[str, Any],
    workflow_data: Dict[str, Any],
    start_time: float
) -> NodeResult:
    """Handler for document_to_text node type
    
    This node converts various document formats to plain text.
    Supported formats: PDF, DOCX, DOC, TXT, RTF, PPTX, etc.
    """
    logger.info(f"Executing Document to Text node {node_id}")
    
    # Extract parameters
    extraction_mode = node_data.get("params", {}).get("extractionMode", "full")  # full, summary, first_page
    include_metadata = node_data.get("params", {}).get("includeMetadata", True)
    preserve_formatting = node_data.get("params", {}).get("preserveFormatting", False)
    max_length = node_data.get("params", {}).get("maxLength", 0)  # 0 means no limit
    variable_name = node_data.get("params", {}).get("variableName", f"document_text_{node_id[:4]}")
    highlighted_text = node_data.get("params", {}).get("highlightedText", [])  # List of highlighted text selections
    
    # Get input data
    input_data = inputs.get("input", {})
    
    if not input_data:
        return NodeResult(
            output={"error": "No input data provided"},
            type="object",
            execution_time=datetime.now().timestamp() - start_time,
            status="error",
            error="No input data provided",
            node_id=node_id,
            node_name=node_data.get("params", {}).get("nodeName", "Document to Text")
        )
    
    try:
        # Determine if input is file data or just file path
        file_content = None
        file_path = None
        file_metadata = {}
        
        if isinstance(input_data, dict):
            if "content" in input_data:
                file_content = input_data.get("content")
                file_metadata = input_data.get("metadata", {})
            elif "path" in input_data:
                file_path = input_data.get("path")
                file_metadata = input_data.get("metadata", {}) or {}
        elif isinstance(input_data, str):
            # Check if it's a base64 string or a file path
            if os.path.exists(input_data):
                file_path = input_data
            else:
                # Assume it's content
                file_content = input_data
                
        if not file_content and not file_path:
            return NodeResult(
                output={"error": "No valid document content or path found in input"},
                type="object",
                execution_time=datetime.now().timestamp() - start_time,
                status="error",
                error="No valid document content or path found in input",
                node_id=node_id,
                node_name=node_data.get("params", {}).get("nodeName", "Document to Text")
            )
        
        # Extract text based on file type
        extracted_text = ""
        metadata = {}
        
        # Get file extension and type
        if file_path:
            filename = os.path.basename(file_path)
            file_extension = os.path.splitext(filename)[1].lower()
            mime_type, _ = mimetypes.guess_type(filename)
        else:
            filename = file_metadata.get("filename", "unknown")
            file_extension = os.path.splitext(filename)[1].lower() if filename != "unknown" else ""
            mime_type = file_metadata.get("type", "")
        
        # Process file content or path
        if not DEPENDENCIES_LOADED:
            return NodeResult(
                output={"error": "Document conversion dependencies not installed"},
                type="object",
                execution_time=datetime.now().timestamp() - start_time,
                status="error",
                error="Document conversion dependencies not installed",
                node_id=node_id,
                node_name=node_data.get("params", {}).get("nodeName", "Document to Text")
            )
        
        # Create a temporary file if content is provided
        temp_file = None
        try:
            if file_content and not file_path:
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
                
                # Write content to temp file
                if isinstance(file_content, str):
                    temp_file.write(file_content.encode('utf-8'))
                else:  # Assume bytes
                    temp_file.write(file_content)
                    
                temp_file.close()
                file_path = temp_file.name
            
            # Extract text based on file type
            if file_extension in ['.pdf']:
                # Extract text from PDF
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)
                    metadata["page_count"] = page_count
                    
                    if extraction_mode == "first_page":
                        # Only extract first page
                        extracted_text = pdf_reader.pages[0].extract_text()
                    else:
                        # Extract all pages
                        for i in range(page_count):
                            page_text = pdf_reader.pages[i].extract_text()
                            if preserve_formatting:
                                extracted_text += f"--- Page {i+1} ---\n{page_text}\n\n"
                            else:
                                extracted_text += f"{page_text} "
            
            elif file_extension in ['.docx']:
                # Extract text from DOCX
                doc = Document(file_path)
                metadata["paragraph_count"] = len(doc.paragraphs)
                
                if preserve_formatting:
                    # With formatting (paragraphs separated by new lines)
                    extracted_text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    # Without formatting (continuous text)
                    extracted_text = " ".join([para.text for para in doc.paragraphs])
            
            else:
                # Use textract for other file types (TXT, RTF, DOC, PPTX, etc.)
                extracted_text = textract.process(file_path).decode('utf-8')
                
                # Clean up text
                if not preserve_formatting:
                    # Replace multiple newlines with space
                    import re
                    extracted_text = re.sub(r'\n+', ' ', extracted_text)
            
            # Apply length limitation if needed
            if max_length > 0 and len(extracted_text) > max_length:
                extracted_text = extracted_text[:max_length] + "..."
            
            # Process highlighted text if provided
            highlighted_content = []
            if highlighted_text:
                for highlight in highlighted_text:
                    start_idx = highlight.get("start")
                    end_idx = highlight.get("end")
                    if start_idx is not None and end_idx is not None and start_idx < len(extracted_text) and end_idx <= len(extracted_text):
                        highlighted_content.append({
                            "text": extracted_text[start_idx:end_idx],
                            "start": start_idx,
                            "end": end_idx
                        })
            
            # Prepare result
            result = {
                "text": extracted_text,
                "highlighted": highlighted_content if highlighted_text else []
            }
            
            # Include metadata if requested
            if include_metadata:
                result["metadata"] = {
                    "filename": filename,
                    "extension": file_extension,
                    "mime_type": mime_type,
                    "size_bytes": os.path.getsize(file_path),
                    **metadata
                }
            
            # Store result in workflow data for variable access
            workflow_data[variable_name] = result
            
            return NodeResult(
                output=result,
                type="object",
                execution_time=datetime.now().timestamp() - start_time,
                status="success",
                node_id=node_id,
                node_name=node_data.get("params", {}).get("nodeName", "Document to Text")
            )
            
        finally:
            # Clean up temp file if created
            if temp_file and os.path.exists(temp_file.name):
                os.unlink(temp_file.name)
    
    except Exception as e:
        logger.error(f"Error in Document to Text node: {str(e)}", exc_info=True)
        return NodeResult(
            output={"error": str(e)},
            type="object",
            execution_time=datetime.now().timestamp() - start_time,
            status="error",
            error=str(e),
            node_id=node_id,
            node_name=node_data.get("params", {}).get("nodeName", "Document to Text")
        ) 